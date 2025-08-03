import sqlite3
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
from collections import Counter

class RoadDamageReportGenerator:
    def __init__(self, db_path="road_damage.db"):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path)
        self.images_dir = "detection_images"
        os.makedirs(self.images_dir, exist_ok=True)
        
    def create_report(self, output_path="road_damage_report.docx"):
        """Generate a comprehensive Word report with images"""
        doc = Document()
        
        # Set up document styles
        self._setup_styles(doc)
        
        # Add title page
        self._add_title_page(doc)
        
        # Add executive summary
        self._add_executive_summary(doc)
        
        # Add detailed findings
        self._add_detailed_findings(doc)
        
        # Add recommendations
        self._add_recommendations(doc)
        
        # Add appendix with images
        self._add_image_appendix(doc)
        
        # Save the document
        doc.save(output_path)
        print(f"Report saved to {output_path}")
        
        return output_path
    
    def _setup_styles(self, doc):
        """Set up custom styles for the report"""
        # Title style
        styles = doc.styles
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading style
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_after = Pt(12)
        
        # Subheading style
        subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_after = Pt(8)
        
        # Body text style
        body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(8)
        body_style.paragraph_format.line_spacing = 1.15
        
        # Caption style
        caption_style = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = 'Arial'
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_title_page(self, doc):
        """Add title page to the document"""
        # Add title
        title = doc.add_paragraph("ROAD DAMAGE DETECTION REPORT", style='CustomTitle')
        title.add_run().add_break()
        
        # Add subtitle
        subtitle = doc.add_paragraph("Automated Analysis System", style='CustomBody')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add report date
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", style='CustomBody')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
    
    def _add_executive_summary(self, doc):
        """Add executive summary section"""
        # Get all detections
        cursor = self.conn.cursor()
        cursor.execute('SELECT * FROM detections ORDER BY timestamp DESC')
        detections = cursor.fetchall()
        
        # Add heading
        heading = doc.add_heading("Executive Summary", level=1)
        heading.style = doc.styles['CustomHeading']
        
        # Summary statistics
        total_detections = len(detections)
        if total_detections > 0:
            damage_types = [d[2] for d in detections]
            pothole_count = damage_types.count('pothole')
            edge_count = damage_types.count('broken_edge')
            
            # Add summary paragraph
            summary = doc.add_paragraph(style='CustomBody')
            summary.add_run("This report presents the findings from the automated road damage detection system. ").bold = True
            summary.add_run(f"A total of {total_detections} road damages were detected, ")
            summary.add_run(f"including {pothole_count} potholes and {edge_count} broken edges. ")
            summary.add_run("The analysis provides severity assessments and maintenance recommendations for each detected damage.")
            
            # Add severity distribution
            severities = []
            for d in detections:
                if d[6]:  # recommendations column
                    try:
                        rec = json.loads(d[6])
                        if isinstance(rec, dict):
                            severities.append(rec.get('severity', 'unknown'))
                    except:
                        pass
            
            if severities:
                severity_counts = {
                    'low': severities.count('low'),
                    'medium': severities.count('medium'),
                    'high': severities.count('high')
                }
                
                severity_para = doc.add_paragraph(style='CustomBody')
                severity_para.add_run("Severity Distribution: ").bold = True
                severity_para.add_run(f"Low: {severity_counts['low']}, ")
                severity_para.add_run(f"Medium: {severity_counts['medium']}, ")
                severity_para.add_run(f"High: {severity_counts['high']}")
        else:
            doc.add_paragraph("No road damages were detected during the analysis period.", style='CustomBody')
        
        doc.add_page_break()
    
    def _add_detailed_findings(self, doc):
        """Add detailed findings section"""
        # Add heading
        heading = doc.add_heading("Detailed Findings", level=1)
        heading.style = doc.styles['CustomHeading']
        
        # Get all detections
        cursor = self.conn.cursor()
        cursor.execute('SELECT * FROM detections ORDER BY timestamp DESC')
        detections = cursor.fetchall()
        
        # Group by damage type
        potholes = [d for d in detections if d[2] == 'pothole']
        broken_edges = [d for d in detections if d[2] == 'broken_edge']
        
        # Add potholes section
        if potholes:
            subheading = doc.add_heading("Potholes", level=2)
            subheading.style = doc.styles['CustomSubheading']
            self._add_damage_table(doc, potholes)
        
        # Add broken edges section
        if broken_edges:
            subheading = doc.add_heading("Broken Edges", level=2)
            subheading.style = doc.styles['CustomSubheading']
            self._add_damage_table(doc, broken_edges)
        
        doc.add_page_break()
    
    def _add_damage_table(self, doc, damages):
        """Add a table of damages"""
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set column widths
        for col in range(5):
            table.columns[col].width = Inches(1.2)
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Date/Time'
        hdr_cells[2].text = 'Severity'
        hdr_cells[3].text = 'Confidence'
        hdr_cells[4].text = 'Recommendation'
        
        # Apply header formatting
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for damage in damages:
            row_cells = table.add_row().cells
            row_cells[0].text = str(damage[1])  # track_id
            row_cells[1].text = damage[5][:16]  # timestamp (truncated)
            
            # Parse severity
            severity = 'N/A'
            if damage[6]:  # recommendations
                try:
                    rec = json.loads(damage[6])
                    if isinstance(rec, dict):
                        severity = rec.get('severity', 'N/A')
                except:
                    pass
            row_cells[2].text = severity.capitalize()
            
            # Format confidence
            row_cells[3].text = f"{damage[3]:.2f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse recommendation
            recommendation = 'N/A'
            if damage[6]:  # recommendations
                try:
                    rec = json.loads(damage[6])
                    if isinstance(rec, dict):
                        recommendation = rec.get('recommendation', 'N/A')
                except:
                    pass
            row_cells[4].text = recommendation
    
    def _add_recommendations(self, doc):
        """Add recommendations section"""
        # Add heading
        heading = doc.add_heading("Maintenance Recommendations", level=1)
        heading.style = doc.styles['CustomHeading']
        
        # Get all recommendations
        cursor = self.conn.cursor()
        cursor.execute('SELECT recommendations FROM detections WHERE recommendations IS NOT NULL')
        recommendations = cursor.fetchall()
        
        # Collect all recommendations
        all_recs = []
        for rec in recommendations:
            try:
                data = json.loads(rec[0])
                if isinstance(data, dict) and 'recommendation' in data:
                    all_recs.append(data['recommendation'])
            except:
                pass
        
        # Count recommendations
        rec_counts = Counter(all_recs)
        top_recs = rec_counts.most_common(5)
        
        # Add recommendations
        doc.add_paragraph("Based on the analysis, the following maintenance actions are recommended:", style='CustomBody')
        
        for i, (rec, count) in enumerate(top_recs, 1):
            p = doc.add_paragraph(style='CustomBody')
            p.add_run(f"{i}. ").bold = True
            p.add_run(rec)
            p.add_run(f" (mentioned {count} times)")
        
        doc.add_page_break()
    
    def _add_image_appendix(self, doc):
        """Add appendix with images of detected damages"""
        # Add heading
        heading = doc.add_heading("Appendix: Damage Images", level=1)
        heading.style = doc.styles['CustomHeading']
        
        # Get all detections with images
        cursor = self.conn.cursor()
        cursor.execute('SELECT track_id, damage_type, recommendations, image_path FROM detections WHERE image_path IS NOT NULL ORDER BY timestamp DESC')
        detections = cursor.fetchall()
        
        if not detections:
            doc.add_paragraph("No images available for detected damages.", style='CustomBody')
            return
        
        # Add images
        for detection in detections:
            track_id, damage_type, recommendations, image_path = detection
            
            # Add subheading
            subheading = doc.add_heading(f"{damage_type.capitalize()} #{track_id}", level=2)
            subheading.style = doc.styles['CustomSubheading']
            
            # Add image
            if os.path.exists(image_path):
                try:
                    doc.add_picture(image_path, width=Inches(4.0))
                    
                    # Add caption
                    caption = doc.add_paragraph(style='CustomCaption')
                    caption.add_run(f"Figure: {damage_type.capitalize()} damage (ID: {track_id})")
                except Exception as e:
                    doc.add_paragraph(f"Error loading image: {str(e)}", style='CustomBody')
            else:
                doc.add_paragraph("Image not found.", style='CustomBody')
            
            # Add details
            details = doc.add_paragraph(style='CustomBody')
            details.add_run("Details: ").bold = True
            
            # Parse recommendations
            if recommendations:
                try:
                    rec_data = json.loads(recommendations)
                    if isinstance(rec_data, dict):
                        severity = rec_data.get('severity', 'N/A')
                        recommendation = rec_data.get('recommendation', 'N/A')
                        
                        details.add_run(f"Severity: {severity.capitalize()}\n")
                        details.add_run(f"Recommendation: {recommendation}")
                except:
                    details.add_run("No recommendation data available.")
            else:
                details.add_run("No recommendation data available.")
            
            # Add page break after every 3 images
            if detections.index(detection) % 3 == 2 and detections.index(detection) != len(detections) - 1:
                doc.add_page_break()
    
    def print_summary(self):
        """Print a summary of the report to the terminal"""
        print("\n" + "="*60)
        print("ROAD DAMAGE DETECTION REPORT SUMMARY")
        print("="*60)
        
        # Get all detections
        cursor = self.conn.cursor()
        cursor.execute('SELECT * FROM detections ORDER BY timestamp DESC')
        detections = cursor.fetchall()
        
        if not detections:
            print("No road damages detected in the database.")
            print("="*60)
            return
        
        # Basic statistics
        total_detections = len(detections)
        damage_types = [d[2] for d in detections]
        pothole_count = damage_types.count('pothole')
        edge_count = damage_types.count('broken_edge')
        
        print(f"Total Detections: {total_detections}")
        print(f"Potholes: {pothole_count}")
        print(f"Broken Edges: {edge_count}")
        
        # Severity distribution
        severities = []
        for d in detections:
            if d[6]:  # recommendations column
                try:
                    rec = json.loads(d[6])
                    if isinstance(rec, dict):
                        severities.append(rec.get('severity', 'unknown'))
                except:
                    pass
        
        if severities:
            severity_counts = {
                'low': severities.count('low'),
                'medium': severities.count('medium'),
                'high': severities.count('high')
            }
            print("\nSeverity Distribution:")
            print(f"  Low: {severity_counts['low']} ({severity_counts['low']/total_detections*100:.1f}%)")
            print(f"  Medium: {severity_counts['medium']} ({severity_counts['medium']/total_detections*100:.1f}%)")
            print(f"  High: {severity_counts['high']} ({severity_counts['high']/total_detections*100:.1f}%)")
        
        # Top recommendations
        all_recs = []
        for d in detections:
            if d[6]:  # recommendations column
                try:
                    rec = json.loads(d[6])
                    if isinstance(rec, dict) and 'recommendation' in rec:
                        all_recs.append(rec['recommendation'])
                except:
                    pass
        
        if all_recs:
            rec_counts = Counter(all_recs)
            print("\nTop Maintenance Recommendations:")
            for i, (rec, count) in enumerate(rec_counts.most_common(3), 1):
                print(f"  {i}. {rec} (mentioned {count} times)")
        
        # Recent detections
        print("\nRecent Detections:")
        for i, d in enumerate(detections[:5], 1):
            track_id = d[1]
            damage_type = d[2]
            timestamp = d[5][:16]  # Truncate timestamp
            confidence = d[3]
            
            severity = 'N/A'
            recommendation = 'N/A'
            if d[6]:
                try:
                    rec = json.loads(d[6])
                    if isinstance(rec, dict):
                        severity = rec.get('severity', 'N/A')
                        recommendation = rec.get('recommendation', 'N/A')
                except:
                    pass
            
            print(f"  {i}. {damage_type.capitalize()} #{track_id} - {timestamp}")
            print(f"     Severity: {severity.capitalize()}, Confidence: {confidence:.2f}")
            print(f"     Recommendation: {recommendation}")
        
        print("="*60)
    
    def close(self):
        """Close database connection"""
        self.conn.close()

def main():
    # Check if database exists
    db_path = "road_damage.db"
    if not os.path.exists(db_path):
        print(f"Error: Database file '{db_path}' not found!")
        print("Please run the detection system first to generate data.")
        return
    
    # Create report generator
    generator = RoadDamageReportGenerator()
    
    try:
        # Generate the report
        print("Generating Word report...")
        report_path = generator.create_report("road_damage_report.docx")
        
        # Print summary to terminal
        generator.print_summary()
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
    finally:
        generator.close()

if __name__ == "__main__":
    main()